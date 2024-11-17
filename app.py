import spaces
import gradio as gr
import json
import os
from pathlib import Path
import logging
from docling.document_converter import DocumentConverter
from docling.datamodel.base_models import InputFormat, DocumentStream
from docling.datamodel.pipeline_options import PdfPipelineOptions, TableFormerMode
from docling.document_converter import PdfFormatOption
import requests
from urllib.parse import urlparse
from datetime import datetime
import tempfile
from docx import Document
from docx.shared import Inches
import markdown

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def is_valid_url(url):
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc])
    except:
        return False

def markdown_to_docx(markdown_content):
    """Convert markdown content to DOCX format"""
    doc = Document()
    
    # Split content into lines
    lines = markdown_content.split('\n')
    
    for line in lines:
        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # Handle lists
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line[3:], style='List Number')
        # Handle normal text
        elif line.strip():
            doc.add_paragraph(line)
        # Handle empty lines
        else:
            doc.add_paragraph()
    
    return doc

def create_output_files(content, original_name):
    """Create temporary files for different formats and return their paths"""
    files = {}
    
    # Generate base filename
    base_name = Path(original_name).stem
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Create markdown file
    md_path = tempfile.NamedTemporaryFile(delete=False, suffix='.md').name
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(content)
    files['markdown'] = md_path
    
    # Create JSON file
    json_content = {
        "title": original_name,
        "content": content,
        "metadata": {
            "conversion_date": datetime.now().isoformat()
        }
    }
    json_path = tempfile.NamedTemporaryFile(delete=False, suffix='.json').name
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(json_content, f, ensure_ascii=False, indent=2)
    files['json'] = json_path
    
    # Create proper DOCX file
    docx_path = tempfile.NamedTemporaryFile(delete=False, suffix='.docx').name
    doc = markdown_to_docx(content)
    doc.save(docx_path)
    files['docx'] = docx_path
    
    return files

@spaces.GPU()
def process_document(input_type, file_input, url_input, use_gpu, table_mode):
    try:
        logger.debug(f"Processing with input type: {input_type}")
        logger.debug(f"File input: {file_input}")
        
        # Configure pipeline
        pipeline_options = PdfPipelineOptions(do_table_structure=True)
        if table_mode:
            pipeline_options.table_structure_options.mode = TableFormerMode.ACCURATE
        else:
            pipeline_options.table_structure_options.mode = TableFormerMode.FAST
            
        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )
        
        # Handle different input types
        if input_type == "file":
            if file_input is None:
                return None, None, None, None, "Please upload a file"
            source = file_input
            original_name = Path(file_input).name
        elif input_type == "url":
            if not url_input or not is_valid_url(url_input):
                return None, None, None, None, "Please enter a valid URL"
            source = url_input
            original_name = Path(urlparse(url_input).path).name or "url_document"
        else:
            return None, None, None, None, "Invalid input type"
            
        # Convert document
        logger.debug(f"Converting document: {source}")
        result = converter.convert(source)
        
        # Get markdown content
        markdown_content = result.document.export_to_markdown()
        
        # Create output files
        output_files = create_output_files(markdown_content, original_name)
        
        return (
            output_files['markdown'],
            output_files['json'],
            output_files['docx'],
            markdown_content,
            "Conversion completed successfully! Use the download buttons below to get your files."
        )
        
    except Exception as e:
        logger.exception("Error occurred during conversion")
        return None, None, None, None, f"Error during conversion: {str(e)}\nCheck the console for detailed error logs."

# Create title HTML with custom style
title_html = """
<div style="text-align: center; max-width: 800px; margin: 0 auto;">
    <h1 style="color: #2C3E50; font-size: 2.5rem; margin-bottom: 0.5rem;">Professional Document Converter</h1>
    <p style="color: #34495E; font-size: 1.1rem; margin-bottom: 1.5rem;">Convert documents from files or URLs to various formats</p>
</div>
"""

# Create Gradio interface with custom theme
with gr.Blocks(css="footer {display: none}") as demo:
    gr.HTML(title_html)
    
    with gr.Row():
        with gr.Column(scale=1):
            input_type = gr.Radio(
                choices=["file", "url"],
                value="file",
                label="Input Type"
            )
            
            # File input with proper file type handling
            file_input = gr.File(
                label="Upload Document",
                file_types=[".pdf", ".PDF"],
                type="filepath"
            )
            
            # URL input
            url_input = gr.Textbox(
                label="Or Enter URL",
                placeholder="https://arxiv.org/pdf/2408.09869"
            )
            
            # Processing options
            use_gpu = gr.Checkbox(label="Use GPU", value=True)
            table_mode = gr.Checkbox(label="Use Accurate Table Mode (Slower but better)", value=False)
            
            convert_btn = gr.Button("Convert Document", variant="primary")
            
        with gr.Column(scale=2):
            # Status message
            status_message = gr.Markdown("")
            
            # Preview area
            preview = gr.Markdown("", label="Preview")
            
            # Download files
            with gr.Group() as download_group:
                gr.Markdown("### Download Files")
                with gr.Row():
                    markdown_output = gr.File(label="Download Markdown")
                    json_output = gr.File(label="Download JSON")
                    docx_output = gr.File(label="Download DOCX")

    # Define the main conversion event
    convert_btn.click(
        fn=process_document,
        inputs=[input_type, file_input, url_input, use_gpu, table_mode],
        outputs=[markdown_output, json_output, docx_output, preview, status_message]
    )
    
    # Add footer
    footer = """
    <div style="text-align: center; margin: 2rem auto; padding: 1rem; border-top: 1px solid #ddd; max-width: 800px;">
        <div style="margin-bottom: 1rem;">
            <a href="https://www.linkedin.com/in/pejman-ebrahimi-4a60151a7/" target="_blank" style="text-decoration: none; color: #2C3E50; margin: 0 10px;">LinkedIn</a> |
            <a href="https://github.com/arad1367" target="_blank" style="text-decoration: none; color: #2C3E50; margin: 0 10px;">GitHub</a> |
            <a href="https://arad1367.pythonanywhere.com/" target="_blank" style="text-decoration: none; color: #2C3E50; margin: 0 10px;">PhD Defense Demo</a> |
            <a href="https://github.com/DS4SD/docling" target="_blank" style="text-decoration: none; color: #2C3E50; margin: 0 10px;">Docling Project</a>
        </div>
        <p style="color: #7F8C8D; margin-top: 0.5rem;">Made with 💖 by Pejman Ebrahimi</p>
    </div>
    """
    gr.HTML(footer)

# Launch the app
if __name__ == "__main__":
    demo.queue(max_size=5)  # Enable queue for better handling of multiple requests
    demo.launch(
        show_error=True,
        share=False,
        debug=True,
        show_api=False,
        server_name="0.0.0.0"
    )